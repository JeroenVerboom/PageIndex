"""
===============================================================================
Document to Markdown Converters - Strategy Pattern
===============================================================================

PURPOSE:
    Provides multiple strategies for converting documents (DOCX, PDF, etc.) 
    to well-structured Markdown that PageIndex can process effectively.

STRATEGIES:
    1. MarkItDownConverter - Uses Microsoft's MarkItDown library (fast, reliable)
    2. LLMMarkdownConverter - Uses GPT to convert each page (more control, slower)
    3. SimpleDocxConverter  - Basic python-docx extraction (fallback)

USAGE:
    from converters import get_converter, ConversionStrategy
    
    # Use MarkItDown (recommended)
    converter = get_converter(ConversionStrategy.MARKITDOWN)
    markdown = converter.convert("/path/to/document.docx")
    
    # Use LLM-based conversion
    converter = get_converter(ConversionStrategy.LLM, api_key="sk-...")
    markdown = converter.convert("/path/to/document.docx")

===============================================================================
"""

import os
from abc import ABC, abstractmethod
from enum import Enum
from typing import Optional
from pathlib import Path


class ConversionStrategy(Enum):
    """
    Available strategies for document-to-markdown conversion.
    
    MARKITDOWN: Uses Microsoft's MarkItDown library
        - Fast and reliable
        - Supports: DOCX, PDF, PPTX, XLSX, HTML, images, audio
        - Best for: Most documents, especially Office formats
    
    LLM: Uses OpenAI GPT to convert content
        - Preserves semantic structure better
        - Can handle unusual formatting
        - Slower and costs API tokens
        - Best for: Complex documents with non-standard layouts
    
    SIMPLE: Basic extraction using python-docx
        - No external dependencies
        - Limited formatting preservation
        - Best for: Simple DOCX files only
    """
    MARKITDOWN = "markitdown"
    LLM = "llm"
    SIMPLE = "simple"


class BaseConverter(ABC):
    """
    Abstract base class for document converters.
    
    All converters must implement the `convert` method that takes a file path
    and returns Markdown content as a string.
    """
    
    @abstractmethod
    def convert(self, file_path: str) -> str:
        """
        Convert a document to Markdown format.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Markdown content as a string
        """
        pass
    
    @property
    @abstractmethod
    def name(self) -> str:
        """Return the converter name for logging purposes."""
        pass
    
    def save(self, markdown: str, output_path: str) -> None:
        """
        Save markdown content to a file.
        
        Args:
            markdown: The markdown content
            output_path: Where to save the file
        """
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown)


# ═══════════════════════════════════════════════════════════════════════════════
# STRATEGY 1: MarkItDown Converter (Recommended)
# ═══════════════════════════════════════════════════════════════════════════════

class MarkItDownConverter(BaseConverter):
    """
    Uses Microsoft's MarkItDown library for document conversion.
    
    HOW IT WORKS:
    MarkItDown is a lightweight utility that converts various file formats
    to Markdown while preserving document structure (headings, lists, tables).
    
    SUPPORTED FORMATS:
    - Microsoft Office: .docx, .xlsx, .pptx
    - PDF documents
    - Images (with OCR if configured)
    - HTML, CSV, JSON, XML
    - ZIP archives (processes contents)
    
    PROS:
    - Fast execution (no API calls)
    - Good structure preservation
    - Supports many formats
    
    CONS:
    - May not handle very complex layouts
    - No semantic understanding of content
    
    INSTALL:
        pip install markitdown
    """
    
    def __init__(self):
        try:
            from markitdown import MarkItDown
            self._md = MarkItDown()
        except ImportError:
            raise ImportError(
                """
                ❌ MarkItDown is required for this converter.
                
                Install it with:
                    pip install markitdown
                """
            )
    
    @property
    def name(self) -> str:
        return "MarkItDown"
    
    def convert(self, file_path: str) -> str:
        """
        Convert a document using MarkItDown.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Markdown content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        result = self._md.convert(file_path)
        return result.text_content


# ═══════════════════════════════════════════════════════════════════════════════
# STRATEGY 2: LLM-based Converter
# ═══════════════════════════════════════════════════════════════════════════════

class LLMMarkdownConverter(BaseConverter):
    """
    Uses OpenAI/Azure OpenAI GPT to convert documents to Markdown.
    
    HOW IT WORKS:
    1. Extracts raw text from the document
    2. Sends the text to GPT with instructions to reformat as Markdown
    3. Returns the structured Markdown output
    
    WHEN TO USE THIS:
    - Documents with complex or unusual layouts
    - When semantic structure is important
    - When other converters produce poor results
    
    PROS:
    - Excellent at inferring structure from content
    - Can handle ambiguous formatting
    - Produces clean, well-organized Markdown
    
    CONS:
    - Requires API calls (costs money)
    - Slower than local conversion
    - May "hallucinate" or modify content
    
    NOTE ON ACCURACY:
    The LLM is instructed to preserve EXACT content and only reformat.
    However, always verify the output for critical documents.
    
    AZURE OPENAI SUPPORT:
    Pass base_url parameter to connect to Azure OpenAI:
        LLMMarkdownConverter(
            api_key="your-key",
            base_url="https://your-resource.openai.azure.com/openai/v1/",
            model="your-deployment-name"
        )
    """
    
    def __init__(self, api_key: str, model: str = "gpt-4.1", base_url: Optional[str] = None):
        """
        Initialize the LLM converter.
        
        Args:
            api_key: OpenAI or Azure OpenAI API key
            model: Model/deployment name to use (default: gpt-4.1)
            base_url: Optional Azure OpenAI endpoint (e.g., https://your-resource.openai.azure.com/openai/v1/)
        """
        self.api_key = api_key
        self.model = model
        self.base_url = base_url
        
        try:
            import openai
            # Support both regular OpenAI and Azure OpenAI
            if base_url:
                # Azure OpenAI uses base_url parameter
                self._client = openai.OpenAI(api_key=api_key, base_url=base_url)
            else:
                # Regular OpenAI
                self._client = openai.OpenAI(api_key=api_key)
        except ImportError:
            raise ImportError(
                """
                ❌ OpenAI library is required for this converter.
                
                Install it with:
                    pip install openai
                """
            )
    
    @property
    def name(self) -> str:
        return f"LLM ({self.model})"
    
    def _extract_raw_text(self, file_path: str) -> str:
        """
        Extract raw text from a document.
        
        Supports DOCX, PDF, and TXT files.
        """
        ext = Path(file_path).suffix.lower()
        
        if ext == ".docx":
            return self._extract_from_docx(file_path)
        elif ext == ".pdf":
            return self._extract_from_pdf(file_path)
        elif ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")
        
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n\n".join(paragraphs)
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 required: pip install PyPDF2")
        
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = [page.extract_text() for page in reader.pages]
        return "\n\n---\n\n".join(pages)
    
    def convert(self, file_path: str) -> str:
        """
        Convert a document to Markdown using an LLM.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Markdown content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        raw_text = self._extract_raw_text(file_path)
        
        # Prompt designed to preserve exact content while adding structure
        prompt = f"""You are a document formatting expert. Convert the following document content to well-structured Markdown.

CRITICAL RULES:
1. PRESERVE ALL CONTENT EXACTLY - do not add, remove, or modify any text
2. Identify and mark headings with appropriate # levels based on context
3. Format lists properly (- for unordered, 1. for ordered)
4. Preserve any tables using Markdown table syntax
5. Keep paragraphs as separate blocks
6. If content appears to be code, wrap it in code blocks
7. Do not add any commentary, introductions, or conclusions

DOCUMENT CONTENT:
{raw_text}

MARKDOWN OUTPUT:"""

        response = self._client.chat.completions.create(
            model=self.model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        
        return response.choices[0].message.content.strip()


# ═══════════════════════════════════════════════════════════════════════════════
# STRATEGY 3: Simple DOCX Converter (Fallback)
# ═══════════════════════════════════════════════════════════════════════════════

class SimpleDocxConverter(BaseConverter):
    """
    Basic DOCX to Markdown converter using python-docx.
    
    This is the original converter from the starter script.
    It extracts paragraphs and converts heading styles to Markdown.
    
    WHEN TO USE:
    - When other converters are not available
    - For simple DOCX files with clear heading styles
    - As a fallback option
    
    LIMITATIONS:
    - Only supports DOCX files
    - Depends on Word's heading styles being properly applied
    - No table or image support
    """
    
    def __init__(self):
        try:
            from docx import Document
            self._Document = Document
        except ImportError:
            raise ImportError(
                """
                ❌ python-docx is required for this converter.
                
                Install it with:
                    pip install python-docx
                """
            )
    
    @property
    def name(self) -> str:
        return "Simple DOCX"
    
    def convert(self, file_path: str) -> str:
        """
        Convert a DOCX file to Markdown using heading style detection.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Markdown content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.lower().endswith(".docx"):
            raise ValueError("SimpleDocxConverter only supports .docx files")
        
        doc = self._Document(file_path)
        markdown_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                markdown_lines.append("")
                continue
            
            # Check paragraph style for heading detection
            style_name = para.style.name.lower() if para.style else ""
            
            if "heading 1" in style_name:
                markdown_lines.append(f"# {text}")
            elif "heading 2" in style_name:
                markdown_lines.append(f"## {text}")
            elif "heading 3" in style_name:
                markdown_lines.append(f"### {text}")
            elif "heading 4" in style_name:
                markdown_lines.append(f"#### {text}")
            elif "title" in style_name:
                markdown_lines.append(f"# {text}")
            else:
                markdown_lines.append(text)
        
        return "\n".join(markdown_lines)


# ═══════════════════════════════════════════════════════════════════════════════
# Factory Function - Get the Right Converter
# ═══════════════════════════════════════════════════════════════════════════════

def get_converter(
    strategy: ConversionStrategy = ConversionStrategy.MARKITDOWN,
    api_key: Optional[str] = None,
    model: str = "gpt-4.1",
    base_url: Optional[str] = None
) -> BaseConverter:
    """
    Factory function to get the appropriate document converter.
    
    This is the main entry point for using converters. It creates and returns
    the converter instance based on the specified strategy.
    
    Args:
        strategy: Which conversion strategy to use (default: MARKITDOWN)
        api_key: OpenAI/Azure OpenAI API key (required for LLM strategy)
        model: Model/deployment name to use for LLM strategy (default: gpt-4.1)
        base_url: Optional Azure OpenAI endpoint (e.g., https://your-resource.openai.azure.com/openai/v1/)
        
    Returns:
        A converter instance implementing BaseConverter
        
    Example:
        >>> # Using MarkItDown
        >>> converter = get_converter(ConversionStrategy.MARKITDOWN)
        >>> markdown = converter.convert("document.docx")
        
        >>> # Using LLM with Azure OpenAI
        >>> converter = get_converter(
        ...     ConversionStrategy.LLM,
        ...     api_key="your-key",
        ...     base_url="https://your-resource.openai.azure.com/openai/v1/",
        ...     model="your-deployment"
        ... )
    """
    if strategy == ConversionStrategy.MARKITDOWN:
        return MarkItDownConverter()
    
    elif strategy == ConversionStrategy.LLM:
        if not api_key:
            raise ValueError("api_key is required for LLM conversion strategy")
        return LLMMarkdownConverter(api_key=api_key, model=model, base_url=base_url)
    
    elif strategy == ConversionStrategy.SIMPLE:
        return SimpleDocxConverter()
    
    else:
        raise ValueError(f"Unknown strategy: {strategy}")


# ═══════════════════════════════════════════════════════════════════════════════
# Markdown Normalizer - Fix Bold-as-Heading Issues
# ═══════════════════════════════════════════════════════════════════════════════

import re

def normalize_markdown_headings(markdown: str) -> str:
    """
    Post-process markdown to convert bold lines that look like headings
    into proper # markdown headings.
    
    THE PROBLEM:
    Many documents use **bold text** instead of proper heading styles.
    When converted to Markdown, these become:
        **1) Section Title**
    Instead of:
        # 1) Section Title
    
    PageIndex requires proper # headings to build its tree structure.
    
    DETECTION HEURISTICS:
    A bold line is likely a heading if:
    1. It's the only content on that line
    2. It starts with a number (e.g., "1)", "1.", "1 -")
    3. It contains keywords like "Extra:", "Stap", "Sectie"
    4. It's relatively short (< 100 chars)
    
    HEADING LEVEL DETECTION:
    - Lines starting with single digit + ) → H2 (##)
    - Lines starting with letter + ) → H3 (###)
    - Lines starting with "Extra", "Stap" → H2
    - Other detected headings → H2 (safe default)
    
    Args:
        markdown: Raw markdown content
        
    Returns:
        Markdown with normalized headings
    """
    lines = markdown.split('\n')
    normalized_lines = []
    
    # Pattern to match standalone bold lines: **text** or *text*
    bold_line_pattern = re.compile(r'^\*\*(.+?)\*\*$')
    
    # Patterns that suggest a line is a heading
    heading_indicators = [
        # Numbered sections: 1), 2., 3 -
        (re.compile(r'^(\d+)\)'), 2),           # "1)" → ## (section)
        (re.compile(r'^(\d+)\.'), 2),           # "1." → ## (section)
        (re.compile(r'^(\d+)\s*[-–—]'), 2),     # "1 -" → ## (section)
        
        # Lettered subsections: A., a), A)
        (re.compile(r'^[A-Z]\)'), 3),           # "A)" → ### (subsection)
        (re.compile(r'^[A-Z]\.'), 3),           # "A." → ### (subsection)
        (re.compile(r'^[a-z]\)'), 3),           # "a)" → ### (subsection)
        
        # Dutch keywords that often indicate headings
        (re.compile(r'^(Stap|Extra|Sectie|Deel|Fase|Module)\s', re.IGNORECASE), 2),
        
        # English keywords
        (re.compile(r'^(Step|Section|Part|Phase|Module)\s', re.IGNORECASE), 2),
        
        # All caps short lines (likely titles)
        (re.compile(r'^[A-Z\s&\-]{5,50}$'), 2),
    ]
    
    for line in lines:
        stripped = line.strip()
        
        # Check if this is a standalone bold line
        bold_match = bold_line_pattern.match(stripped)
        
        if bold_match:
            inner_text = bold_match.group(1).strip()
            
            # Determine if this looks like a heading
            heading_level = None
            
            for pattern, level in heading_indicators:
                if pattern.search(inner_text):
                    heading_level = level
                    break
            
            # If no pattern matched but it's short, treat as heading
            if heading_level is None and len(inner_text) < 80:
                # Check if it's likely a heading (not a regular emphasis)
                # A heading usually doesn't have periods in the middle
                if not re.search(r'\.\s+[a-z]', inner_text):
                    heading_level = 2  # Default to H2
            
            if heading_level:
                # Convert to proper heading
                prefix = '#' * heading_level
                normalized_lines.append(f"{prefix} {inner_text}")
            else:
                # Keep as bold (probably just emphasis)
                normalized_lines.append(line)
        else:
            # Not a bold line, keep as-is
            normalized_lines.append(line)
    
    result = '\n'.join(normalized_lines)
    
    # Count how many headings we created
    heading_count = len(re.findall(r'^#{1,6}\s+', result, re.MULTILINE))
    print(f"   📝 Normalized markdown: found {heading_count} headings")
    
    return result


# ═══════════════════════════════════════════════════════════════════════════════
# Convenience Function - Convert with Auto-Detection
# ═══════════════════════════════════════════════════════════════════════════════

def convert_document(
    file_path: str,
    output_path: Optional[str] = None,
    strategy: ConversionStrategy = ConversionStrategy.MARKITDOWN,
    api_key: Optional[str] = None,
    model: str = "gpt-4.1",
    normalize_headings: bool = True,
    base_url: Optional[str] = None
) -> str:
    """
    High-level convenience function to convert a document to Markdown.
    
    Args:
        file_path: Path to the input document
        output_path: Optional path to save the Markdown output
        strategy: Conversion strategy to use
        api_key: OpenAI API key (for LLM strategy)
        model: Model to use (for LLM strategy)
        normalize_headings: If True, convert bold lines to proper # headings
        base_url: Optional Azure OpenAI endpoint (for LLM strategy)
        
    Returns:
        Markdown content as a string
        
    Example:
        >>> # Using MarkItDown
        >>> md = convert_document("report.docx", "report.md")
        
        >>> # Using LLM with Azure
        >>> md = convert_document(
        ...     "report.docx",
        ...     strategy=ConversionStrategy.LLM,
        ...     api_key="your-key",
        ...     base_url="https://...",
        ...     model="your-deployment"
        ... )
    """
    converter = get_converter(strategy, api_key, model, base_url)
    print(f"📄 Converting with {converter.name}...")
    
    markdown = converter.convert(file_path)
    
    # Post-process to fix heading issues
    if normalize_headings:
        markdown = normalize_markdown_headings(markdown)
    
    if output_path:
        converter.save(markdown, output_path)
        print(f"✅ Saved to: {output_path}")
    
    return markdown


# ═══════════════════════════════════════════════════════════════════════════════
# Test / Demo
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    # Demo usage
    print("Document Converter Strategies")
    print("=" * 50)
    
    for strategy in ConversionStrategy:
        print(f"  - {strategy.value}: {strategy.name}")
    
    print("\nUsage:")
    print("  from converters import get_converter, ConversionStrategy")
    print("  converter = get_converter(ConversionStrategy.MARKITDOWN)")
    print("  markdown = converter.convert('document.docx')")
